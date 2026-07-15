import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from bs4 import BeautifulSoup, NavigableString

def set_cell_background(cell, fill_hex):
    """Met un arrière-plan coloré sur une cellule Word."""
    tcPr = cell._tc.get_or_add_tcPr()
    shading = parse_xml(r'<w:shd %s w:fill="%s"/>' % (nsdecls('w'), fill_hex))
    tcPr.append(shading)

def generate_docx():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    html_path = os.path.join(base_dir, "memoire_complet.html")
    docx_path = os.path.join(base_dir, "memoire_complet.docx")

    if not os.path.exists(html_path):
        print(f"Erreur : Le fichier HTML n'existe pas dans {html_path}")
        return

    print(f"Lecture du mémoire compilé : {html_path}...")
    with open(html_path, 'r', encoding='utf-8') as f:
        html_content = f.read()

    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Trouver le conteneur principal du mémoire
    container = soup.find('div', class_='memoire-container')
    if not container:
        container = soup.body

    print("Création du document Word académique...")
    doc = Document()

    # Configuration du format A4 et des marges (2.0 cm partout pour optimiser l'espace)
    for section in doc.sections:
        section.page_width = Inches(8.27)   # A4
        section.page_height = Inches(11.69)  # A4
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Style par défaut pour le corps du texte
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Georgia'
    font_normal.size = Pt(11.5)
    font_normal.color.rgb = RGBColor(26, 26, 26) # Gris très foncé académique (#1a1a1a)

    def parse_inline_elements(element, paragraph, default_italic=False):
        """Parcourt les éléments textuels enfants pour appliquer les styles (gras, italique, code)."""
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text.strip() or text == ' ':
                    run = paragraph.add_run(text)
                    run.font.name = 'Georgia'
                    run.font.size = Pt(11.5)
                    if default_italic:
                        run.font.italic = True
            elif child.name in ['strong', 'b']:
                run = paragraph.add_run(child.get_text())
                run.font.name = 'Georgia'
                run.font.size = Pt(11.5)
                run.font.bold = True
                if default_italic:
                    run.font.italic = True
            elif child.name in ['em', 'i']:
                run = paragraph.add_run(child.get_text())
                run.font.name = 'Georgia'
                run.font.size = Pt(11.5)
                run.font.italic = True
            elif child.name == 'code':
                run = paragraph.add_run(child.get_text())
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(120, 20, 60) # Rouge bordeaux discret pour le code
            elif child.name == 'br':
                paragraph.add_run('\n')
            elif child.name == 'a':
                run = paragraph.add_run(child.get_text())
                run.font.name = 'Georgia'
                run.font.size = Pt(11.5)
                run.font.underline = True
                run.font.color.rgb = RGBColor(30, 58, 138) # Bleu lien
            elif child.name in ['span', 'small']:
                # Traitement récursif léger pour les spans imbriqués
                parse_inline_elements(child, paragraph, default_italic)

    # Parcourir et traiter les éléments direct de premier niveau
    children = list(container.children)
    i = 0
    total_elements = len(children)

    # Variables pour éviter les sauts de page doublés
    has_just_added_page_break = False

    while i < total_elements:
        el = children[i]
        i += 1

        if isinstance(el, NavigableString):
            continue

        # 1. Gestion des Titres
        if el.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(el.name[1])
            text = el.get_text().strip()

            if not text:
                continue

            # Détecter les sections nécessitant un saut de page (standards académiques)
            needs_page_break = False
            text_upper = text.upper()
            if level <= 2:
                if any(x in text_upper for x in ["CHAPITRE", "INTRODUCTION", "CONCLUSION", "BIBLIOGRAPHIE", "REMERCIEMENTS", "CONTENU"]):
                    needs_page_break = True

            if needs_page_break and not has_just_added_page_break:
                doc.add_page_break()
                has_just_added_page_break = True
            else:
                has_just_added_page_break = False

            # Créer le titre dans Word
            h = doc.add_heading(level=level)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(8)
            h.paragraph_format.keep_with_next = True

            run = h.add_run(text)
            run.font.name = 'Calibri'
            run.font.bold = True
            run.font.color.rgb = RGBColor(30, 58, 138) # Deep Academic Blue #1e3a8a

            # Formater la taille selon le niveau du titre
            if level == 1:
                run.font.size = Pt(22)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 2:
                run.font.size = Pt(16)
                h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif level == 3:
                run.font.size = Pt(13)
                h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                run.font.size = Pt(11.5)
                h.alignment = WD_ALIGN_PARAGRAPH.LEFT

            continue

        # Réinitialiser le marqueur de saut de page si nous passons à un élément textuel
        has_just_added_page_break = False

        # 2. Gestion des sauts de page explicites du HTML
        if el.name == 'div' and 'page-break' in el.get('class', []):
            doc.add_page_break()
            has_just_added_page_break = True
            continue

        # 3. Gestion des listes à puces ou numérotées
        if el.name in ['ul', 'ol']:
            style_name = 'List Bullet' if el.name == 'ul' else 'List Number'
            for li in el.find_all('li', recursive=False):
                p = doc.add_paragraph(style=style_name)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.15
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                parse_inline_elements(li, p)
            continue

        # 4. Gestion des citations (blockquote)
        if el.name == 'blockquote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing = 1.15
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            parse_inline_elements(el, p, default_italic=True)
            continue

        # 5. Gestion des Blocs de Code source (pre)
        if el.name == 'pre':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing = 1.0
            
            run = p.add_run(el.get_text())
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(64, 64, 64)
            continue

        # 6. Gestion des Tableaux (table)
        if el.name == 'table':
            rows = el.find_all('tr')
            if not rows:
                continue

            max_cols = 0
            for r in rows:
                cols = len(r.find_all(['td', 'th']))
                if cols > max_cols:
                    max_cols = cols

            table = doc.add_table(rows=len(rows), cols=max_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True

            # Ajouter des bordures élégantes via XML
            tblPr = table._tbl.tblPr
            tblBorders = parse_xml(
                r'<w:tblBorders %s>'
                r'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
                r'  <w:left w:val="none"/>'
                r'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
                r'  <w:right w:val="none"/>'
                r'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
                r'  <w:insideV w:val="none"/>'
                r'</w:tblBorders>' % nsdecls('w')
            )
            tblPr.append(tblBorders)

            for row_idx, r in enumerate(rows):
                cells = r.find_all(['td', 'th'])
                for col_idx, cell in enumerate(cells):
                    if col_idx >= max_cols:
                        break
                    w_cell = table.cell(row_idx, col_idx)
                    p = w_cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(3)

                    is_header = cell.name == 'th'
                    if is_header:
                        set_cell_background(w_cell, "F1F5F9") # Couleur d'entête universitaire

                    for child in cell.children:
                        if isinstance(child, NavigableString):
                            run = p.add_run(str(child))
                        elif child.name in ['strong', 'b'] or is_header:
                            run = p.add_run(child.get_text())
                            run.font.bold = True
                        else:
                            run = p.add_run(child.get_text())
                        
                        run.font.name = 'Calibri'
                        run.font.size = Pt(10)
                        if is_header:
                            run.font.color.rgb = RGBColor(30, 58, 138)
            continue

        # 7. Gestion des DIV complexes (screenshot-cards, logo-grids, ou conteneurs d'images)
        if el.name == 'div':
            # Rechercher récursivement des images dans ce DIV
            img_tags = el.find_all('img')
            if img_tags:
                # C'est un conteneur d'images !
                for img_tag in img_tags:
                    src = img_tag.get('src')
                    alt = img_tag.get('alt', 'Figure')
                    if src:
                        # Nettoyer les chemins relatifs
                        src_path = os.path.join(base_dir, src)
                        if os.path.exists(src_path):
                            print(f"Insertion de l'image : {src}...")
                            
                            # Choix de la taille selon le type d'image
                            if "app_" in src:
                                width = Inches(3.2) # Écran d'app mobile vertical élancé
                            elif "dashboard_" in src or "diagramme_" in src:
                                width = Inches(6.2) # Dashboard web ou Diagramme UML large
                            elif "logo_" in src or "flutter_logo" in src:
                                width = Inches(1.5) # Logo discret
                            else:
                                width = Inches(4.5) # Par défaut

                            # Insérer l'image centrée avec haute compatibilité et sans alinéa textuel
                            doc.add_picture(src_path, width=width)
                            p_img = doc.paragraphs[-1]
                            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p_img.paragraph_format.space_before = Pt(12)
                            p_img.paragraph_format.space_after = Pt(6)
                            p_img.paragraph_format.first_line_indent = Inches(0)
                            p_img.paragraph_format.left_indent = Inches(0)
                            p_img.paragraph_format.right_indent = Inches(0)
                            p_img.paragraph_format.line_spacing = 1.0
                            
                            # Recherche de légende ou description associée dans le conteneur
                            # 1. Légende Figure (ex: Figure 4.10 ...)
                            caption_text = ""
                            legend_el = el.find(['h5', 'h6', 'p'], style=lambda s: s and ('font-family' in s or 'font-size' in s))
                            if not legend_el:
                                legend_el = el.find(['h5', 'h6'])
                            if legend_el:
                                caption_text = legend_el.get_text().strip()
                            
                            if caption_text:
                                p_cap = doc.add_paragraph()
                                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                p_cap.paragraph_format.space_after = Pt(12)
                                p_cap.paragraph_format.keep_with_next = True
                                run_cap = p_cap.add_run(caption_text)
                                run_cap.font.name = 'Calibri'
                                run_cap.font.bold = True
                                run_cap.font.size = Pt(10.5)
                                run_cap.font.color.rgb = RGBColor(30, 58, 138)
                            
                            # 2. Paragraphe descriptif accompagnant la figure
                            desc_text = ""
                            desc_el = el.find('p', style=lambda s: s and ('text-align: justify' in s or 'font-family: var(--font-serif)' in s))
                            if desc_el and desc_el.get_text().strip() != caption_text:
                                desc_text = desc_el.get_text().strip()
                            
                            if desc_text:
                                p_desc = doc.add_paragraph()
                                p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                p_desc.paragraph_format.space_after = Pt(24)
                                p_desc.paragraph_format.line_spacing = 1.15
                                run_desc = p_desc.add_run(desc_text)
                                run_desc.font.name = 'Georgia'
                                run_desc.font.size = Pt(10.5)
                                run_desc.font.italic = True
                                run_desc.font.color.rgb = RGBColor(80, 80, 80)
                continue
            
            # Si le DIV ne contient pas d'images mais contient du texte brut (ex: les badges DevOps)
            text_content = el.get_text().strip()
            if text_content:
                # S'il contient des sous-éléments de badges
                badges = el.find_all('div', style=lambda s: s and 'background' in s)
                if badges:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(8)
                    p.paragraph_format.space_after = Pt(12)
                    
                    run_intro = p.add_run("Écosystème DevOps & Services de déploiement : ")
                    run_intro.font.bold = True
                    run_intro.font.size = Pt(11)
                    
                    badge_names = [b.get_text().strip() for b in badges if b.get_text().strip()]
                    run_badges = p.add_run(" | ".join(badge_names))
                    run_badges.font.italic = True
                    run_badges.font.size = Pt(11)
                    run_badges.font.color.rgb = RGBColor(30, 58, 138)
                else:
                    # Div textuel standard
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.space_after = Pt(12)
                    parse_inline_elements(el, p)
            continue

        # 8. Gestion des Paragraphes normaux (p)
        if el.name == 'p':
            img_tag = el.find('img')
            if img_tag:
                # Image simple dans un paragraphe
                src = img_tag.get('src')
                if src:
                    src_path = os.path.join(base_dir, src)
                    if os.path.exists(src_path):
                        print(f"Insertion de l'image (simple p) : {src}...")
                        width = Inches(6.2) if ("diagramme_" in src or "dashboard_" in src) else Inches(3.2)
                        
                        # Insérer l'image simple centrée avec haute compatibilité et sans alinéa
                        doc.add_picture(src_path, width=width)
                        p_img = doc.paragraphs[-1]
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_img.paragraph_format.space_before = Pt(12)
                        p_img.paragraph_format.space_after = Pt(12)
                        p_img.paragraph_format.first_line_indent = Inches(0)
                        p_img.paragraph_format.left_indent = Inches(0)
                        p_img.paragraph_format.right_indent = Inches(0)
                        p_img.paragraph_format.line_spacing = 1.0
                continue

            # Paragraphe de texte standard
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(12)
            
            # Appliquer le retrait de 0.5 pouce pour la première ligne du paragraphe académique
            # sauf si c'est un paragraphe de description ou de légende de figure ou s'il a la classe "no-indent"
            classes = el.get('class', [])
            if 'no-indent' not in classes and not el.get_text().strip().startswith(('Figure', 'Source', 'Légende')):
                p.paragraph_format.first_line_indent = Inches(0.5)

            parse_inline_elements(el, p)

    print(f"Sauvegarde du fichier Word final dans : {docx_path}...")
    doc.save(docx_path)
    print("Document Word généré avec un succès absolu !")

if __name__ == "__main__":
    generate_docx()
